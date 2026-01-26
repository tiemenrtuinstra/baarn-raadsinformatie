# -*- coding: utf-8 -*-
import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Laad alle documenten
with open('data/soestdijk_alle_docs.json', 'r', encoding='utf-8') as f:
    all_docs = json.load(f)

# Laad moties/amendementen (inclusief zonder inhoud)
with open('data/soestdijk_moties_amendementen.json', 'r', encoding='utf-8') as f:
    moties_amen = json.load(f)

doc = Document()

# Stel standaard lettertype in
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(9)

# Titel
title = doc.add_heading('PALEIS SOESTDIJK', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('COMPLETE DOCUMENTATIE')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(16)
subtitle.runs[0].bold = True

doc.add_paragraph('Alle Moties, Amendementen, Inspraakreacties, Vragen en Standpunten')
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Gemeente Baarn - Raadsinformatie 2011-2026')
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_paragraph(f'Dit document bevat {len(all_docs)} documenten met volledige tekst.')
doc.add_page_break()

# =============================================================================
# DEEL 1: MOTIES EN AMENDEMENTEN
# =============================================================================
doc.add_heading('DEEL 1: MOTIES EN AMENDEMENTEN', level=1)

doc.add_paragraph(f'Totaal: {len(moties_amen)} moties en amendementen')
doc.add_paragraph()

# Sorteer op datum
moties_amen_sorted = sorted(moties_amen, key=lambda x: x['date'] if x['date'] != 'onbekend' else '9999')

for item in moties_amen_sorted:
    # Skip bijlagen
    if 'bijlage' in item['title'].lower() and ('regels' in item['title'].lower() or 'toelichting' in item['title'].lower() or 'verbeelding' in item['title'].lower()):
        continue

    doc.add_heading(f"{item['type']} - {item['date']}", level=2)

    p = doc.add_paragraph()
    p.add_run('Titel: ').bold = True
    p.add_run(item['title'])

    if item['partijen']:
        p = doc.add_paragraph()
        p.add_run('Indieners: ').bold = True
        p.add_run(', '.join(item['partijen']))

    p = doc.add_paragraph()
    p.add_run('Status: ').bold = True
    p.add_run(item['status'])

    doc.add_paragraph()

    if item['content'] and len(item['content']) > 50:
        doc.add_paragraph('VOLLEDIGE TEKST:').runs[0].bold = True
        content = item['content'][:10000].replace('\x00', '').replace('\r', '')
        doc.add_paragraph(content)
    else:
        doc.add_paragraph('(Geen tekstinhoud beschikbaar in database)')

    doc.add_paragraph('_' * 80)
    doc.add_paragraph()

doc.add_page_break()

# =============================================================================
# DEEL 2: SCHRIFTELIJKE EN MONDELINGE VRAGEN
# =============================================================================
doc.add_heading('DEEL 2: SCHRIFTELIJKE EN MONDELINGE VRAGEN', level=1)

vragen = [d for d in all_docs if d['category'] in ['Schriftelijke vragen', 'Mondelinge vragen']]
vragen_sorted = sorted(vragen, key=lambda x: x['date'] if x['date'] != 'onbekend' else '9999')

doc.add_paragraph(f'Totaal: {len(vragen)} documenten')
doc.add_paragraph()

for v in vragen_sorted:
    doc.add_heading(f"{v['category']} - {v['date']}", level=2)

    p = doc.add_paragraph()
    p.add_run('Titel: ').bold = True
    p.add_run(v['title'])

    doc.add_paragraph()
    doc.add_paragraph('VOLLEDIGE TEKST:').runs[0].bold = True
    content = v['content'][:8000].replace('\x00', '').replace('\r', '')
    doc.add_paragraph(content)

    doc.add_paragraph('_' * 80)
    doc.add_paragraph()

doc.add_page_break()

# =============================================================================
# DEEL 3: INSPRAAKREACTIES EN ZIENSWIJZEN
# =============================================================================
doc.add_heading('DEEL 3: INSPRAAKREACTIES EN ZIENSWIJZEN', level=1)

inspraak = [d for d in all_docs if d['category'] in ['Inspraak', 'Zienswijze']]
inspraak_sorted = sorted(inspraak, key=lambda x: x['date'] if x['date'] != 'onbekend' else '9999')

doc.add_paragraph(f'Totaal: {len(inspraak)} documenten met tekst')
doc.add_paragraph()

for ins in inspraak_sorted:
    doc.add_heading(f"Inspraak - {ins['date']}", level=2)

    p = doc.add_paragraph()
    p.add_run('Titel: ').bold = True
    p.add_run(ins['title'])

    doc.add_paragraph()
    doc.add_paragraph('VOLLEDIGE TEKST:').runs[0].bold = True
    content = ins['content'][:10000].replace('\x00', '').replace('\r', '')
    doc.add_paragraph(content)

    doc.add_paragraph('_' * 80)
    doc.add_paragraph()

doc.add_page_break()

# =============================================================================
# DEEL 4: BRIEVEN EN REACTIES EXTERNE PARTIJEN
# =============================================================================
doc.add_heading('DEEL 4: BRIEVEN EN REACTIES EXTERNE PARTIJEN', level=1)

brieven = [d for d in all_docs if d['category'] in ['Brief', 'Reactie', 'Stichting Parel']]
brieven_sorted = sorted(brieven, key=lambda x: x['date'] if x['date'] != 'onbekend' else '9999')

doc.add_paragraph(f'Totaal: {len(brieven)} documenten')
doc.add_paragraph()

for br in brieven_sorted:
    doc.add_heading(f"{br['category']} - {br['date']}", level=2)

    p = doc.add_paragraph()
    p.add_run('Titel: ').bold = True
    p.add_run(br['title'])

    doc.add_paragraph()
    doc.add_paragraph('VOLLEDIGE TEKST:').runs[0].bold = True
    content = br['content'][:10000].replace('\x00', '').replace('\r', '')
    doc.add_paragraph(content)

    doc.add_paragraph('_' * 80)
    doc.add_paragraph()

doc.add_page_break()

# =============================================================================
# DEEL 5: JURIDISCHE ADVIEZEN (HBR ADVOCATEN)
# =============================================================================
doc.add_heading('DEEL 5: JURIDISCHE ADVIEZEN', level=1)

juridisch = [d for d in all_docs if d['category'] == 'Juridisch advies']
juridisch_sorted = sorted(juridisch, key=lambda x: x['date'] if x['date'] != 'onbekend' else '9999')

doc.add_paragraph(f'Totaal: {len(juridisch)} documenten')
doc.add_paragraph()

for jur in juridisch_sorted:
    doc.add_heading(f"Juridisch advies - {jur['date']}", level=2)

    p = doc.add_paragraph()
    p.add_run('Titel: ').bold = True
    p.add_run(jur['title'])

    doc.add_paragraph()
    doc.add_paragraph('VOLLEDIGE TEKST:').runs[0].bold = True
    content = jur['content'][:12000].replace('\x00', '').replace('\r', '')
    doc.add_paragraph(content)

    doc.add_paragraph('_' * 80)
    doc.add_paragraph()

doc.add_page_break()

# =============================================================================
# DEEL 6: COA DOCUMENTEN
# =============================================================================
doc.add_heading('DEEL 6: COA OPVANG DOCUMENTEN', level=1)

coa = [d for d in all_docs if d['category'] == 'COA' or 'coa' in d['title'].lower() or 'asiel' in d['title'].lower()]
coa_sorted = sorted(coa, key=lambda x: x['date'] if x['date'] != 'onbekend' else '9999')

doc.add_paragraph(f'Totaal: {len(coa)} documenten')
doc.add_paragraph()

for c in coa_sorted:
    doc.add_heading(f"COA - {c['date']}", level=2)

    p = doc.add_paragraph()
    p.add_run('Titel: ').bold = True
    p.add_run(c['title'])

    doc.add_paragraph()
    doc.add_paragraph('VOLLEDIGE TEKST:').runs[0].bold = True
    content = c['content'][:10000].replace('\x00', '').replace('\r', '')
    doc.add_paragraph(content)

    doc.add_paragraph('_' * 80)
    doc.add_paragraph()

doc.add_page_break()

# =============================================================================
# DEEL 7: RAADSINFORMATIEBRIEVEN
# =============================================================================
doc.add_heading('DEEL 7: RAADSINFORMATIEBRIEVEN', level=1)

ribs = [d for d in all_docs if d['category'] == 'Raadsinformatiebrief']
ribs_sorted = sorted(ribs, key=lambda x: x['date'] if x['date'] != 'onbekend' else '9999')

doc.add_paragraph(f'Totaal: {len(ribs)} documenten')
doc.add_paragraph()

for rib in ribs_sorted:
    doc.add_heading(f"RIB - {rib['date']}", level=2)

    p = doc.add_paragraph()
    p.add_run('Titel: ').bold = True
    p.add_run(rib['title'])

    doc.add_paragraph()
    doc.add_paragraph('VOLLEDIGE TEKST:').runs[0].bold = True
    content = rib['content'][:12000].replace('\x00', '').replace('\r', '')
    doc.add_paragraph(content)

    doc.add_paragraph('_' * 80)
    doc.add_paragraph()

# Afsluiting
doc.add_page_break()
doc.add_heading('BRONVERMELDING', level=1)

doc.add_paragraph('''Dit document is samengesteld op basis van de Baarn Raadsinformatie database.

Database statistieken:
- 659 vergaderingen
- 7.068 agenda items
- 16.763 documenten
- Periode: december 2009 - januari 2026

Specifiek voor Paleis Soestdijk:
- 44 agenda items
- 295 documenten totaal
- 113 documenten met tekstinhoud
- 28 moties en amendementen
- 15 schriftelijke/mondelinge vragen
- 60 inspraakreacties (titels)''')

doc.add_paragraph()
p = doc.add_paragraph('Document gegenereerd: januari 2026')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph('Bron: Baarn Raadsinformatie MCP Server')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Opslaan
output_path = 'data/Paleis_Soestdijk_COMPLEET_alle_teksten.docx'
doc.save(output_path)
print(f'Document opgeslagen: {output_path}')
