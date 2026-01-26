# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# Stel standaard lettertype in
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Titel
title = doc.add_heading('PALEIS SOESTDIJK', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Volledige Tijdlijn, Standpunten en COA-opvang')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.runs[0]
subtitle_run.font.size = Pt(14)
subtitle_run.font.italic = True

doc.add_paragraph('Gemeente Baarn - Raadsinformatie 2011-2026')
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Inhoudsopgave
doc.add_heading('INHOUDSOPGAVE', level=1)
toc_items = [
    '1. Samenvatting en Kerngegevens',
    '2. Historische Achtergrond',
    '3. Chronologische Tijdlijn 2011-2026',
    '4. Standpunten Politieke Partijen',
    '5. Standpunten Externe Organisaties',
    '6. COA-opvang op het Marechausseeterrein',
    '7. Raad van State Uitspraak en Gevolgen',
    '8. Huidige Stand van Zaken (2025-2026)',
    '9. Bronnen en Documenten'
]
for item in toc_items:
    doc.add_paragraph(item)

doc.add_page_break()

# SECTIE 1: Samenvatting
doc.add_heading('1. SAMENVATTING EN KERNGEGEVENS', level=1)

doc.add_heading('Statistieken', level=2)
table = doc.add_table(rows=5, cols=2)
table.style = 'Table Grid'
data = [
    ('Gegeven', 'Aantal'),
    ('Agenda items', '44'),
    ('Documenten', '295'),
    ('Betrokken vergaderingen', '659'),
    ('Periode', 'December 2009 - Januari 2026')
]
for i, (col1, col2) in enumerate(data):
    table.rows[i].cells[0].text = col1
    table.rows[i].cells[1].text = col2

doc.add_paragraph()

doc.add_heading('Kernbesluiten', level=2)
besluiten = [
    ('2012', 'Omgevingsvisie Paleis Soestdijk vastgesteld'),
    ('2019', 'Ruimtelijk Kader Landgoed Paleis Soestdijk vastgesteld'),
    ('2020', 'Voorontwerpbestemmingsplan vastgesteld'),
    ('2021', 'Ontwerpbestemmingsplan vastgesteld (met amendementen)'),
    ('23 februari 2022', 'BESTEMMINGSPLAN DEFINITIEF VASTGESTELD'),
    ('Januari 2024', 'Raad van State vernietigt delen bestemmingsplan'),
    ('September 2024', 'Reflectiedocument besluitvormingsproces vastgesteld'),
    ('2025-2026', 'Vervolgproces loopt, COA-discussie actueel')
]
for jaar, besluit in besluiten:
    p = doc.add_paragraph()
    p.add_run(f'{jaar}: ').bold = True
    p.add_run(besluit)

doc.add_page_break()

# SECTIE 2: Historische Achtergrond
doc.add_heading('2. HISTORISCHE ACHTERGROND', level=1)

doc.add_heading('Eigendomsgeschiedenis', level=2)
doc.add_paragraph('''In 1971 werd het Rijk eigenaar van Paleis Soestdijk. Na het overlijden van Prinses Juliana in 2004 en Prins Bernhard in 2004 werd het paleis niet meer bewoond door leden van het Koninklijk Huis.

Om tot herbestemming te komen heeft de "Ronde Tafel Paleis Soestdijk" in juli 2015 advies uitgebracht. Vervolgens is het paleis middels een verkoopprocedure in de markt gezet.''')

doc.add_heading('Verkoopprocedure en Made by Holland', level=2)
doc.add_paragraph('''Voor de verkoopprocedure zijn uitgangspunten meegegeven waar elke inschrijving aan moest voldoen. Er is bewust aansluiting gezocht bij de "Omgevingsvisie Paleis Soestdijk" (2011). Het belangrijkste uitgangspunt was dat het ensemble van paleis, park en bos als geheel behouden moest blijven voor de toekomst.

Het plan "Made by Holland" van de MeyerBergman Erfgoed Groep (MBEG) is na een selectieprocedure en biedingsfase door de rijksoverheid tot winnaar uitgeroepen. Geen van de drie partijen die uitgenodigd waren een bieding te doen had ontbindende voorwaarden verbonden aan de bieding.

Eind 2017 werd MBEG eigenaar van het landgoed via:
- Koopovereenkomst: 3 juli 2017
- Akte van levering: 20 december 2017''')

doc.add_heading('Het Plan Made by Holland', level=2)
doc.add_paragraph('''Het plan voorzag in:
1. Restauratie van het paleis en de tuinen
2. Woningbouw in het Alexanderkwartier (op locatie voormalige kazerne) om de restauratie te bekostigen
3. Behoud van het Borrebos en natuurgebied
4. Publieke toegankelijkheid van het landgoed
5. Culturele en zakelijke evenementen in het paleis''')

doc.add_page_break()

# SECTIE 3: Chronologische Tijdlijn
doc.add_heading('3. CHRONOLOGISCHE TIJDLIJN 2011-2026', level=1)

# 2011-2012
doc.add_heading('2011-2012: Start van het Dossier', level=2)
doc.add_paragraph('26 oktober 2011: Eerste motie BOP over Paleis Soestdijk')
doc.add_paragraph('15 februari 2012: Debat Omgevingsvisie Paleis Soestdijk')
doc.add_paragraph('29 februari 2012: VASTSTELLING OMGEVINGSVISIE').runs[0].bold = True

p = doc.add_paragraph()
p.add_run('Inhoud Omgevingsvisie: ').bold = True
p.add_run('De omgevingsvisie schetst het wensbeeld voor het paleis en omgeving. Op initiatief van de provincie is de visie opgesteld door gemeente Baarn, gemeente Soest en provincie Utrecht. Centraal staat de wens dat het paleis met directe omgeving een poortfunctie krijgt voor de Heuvelrug.')

# 2016
doc.add_heading('2016: Herbestemming', level=2)
doc.add_paragraph('13 januari 2016: RIB Ontwikkeling Paleis Soestdijk')
doc.add_paragraph('9 november 2016: RIB vervolg herbestemming landgoed en paleis Soestdijk')

p = doc.add_paragraph()
p.add_run('Externe reacties: ').bold = True
doc.add_paragraph('- Open brief aan gemeenteraad m.b.t. Paleis Soestdijk', style='List Bullet')
doc.add_paragraph('- Stichting Mooi Baarn: de vier plannen voor Paleis Soestdijk', style='List Bullet')

# 2019
doc.add_heading('2019: Ruimtelijk Kader', level=2)
doc.add_paragraph('3 april 2019: Informatie Ruimtelijk Kader')
doc.add_paragraph('10 april 2019: Debat Ruimtelijk Kader')
doc.add_paragraph('17 april 2019: VASTSTELLING RUIMTELIJK KADER').runs[0].bold = True

doc.add_heading('Moties april 2019:', level=3)
doc.add_paragraph('- Motie Randvoorwaarden voor het opstellen bestemmingsplan (AANGENOMEN)', style='List Bullet')
doc.add_paragraph('- Motie Opstellen van een gebiedsvisie (AANGENOMEN)', style='List Bullet')

doc.add_heading('Inspraakreacties april 2019:', level=3)
insprekers = [
    'Stichting de Parel van Baarn - Open brief over behoud historisch ensemble',
    'Natuur en Milieufederatie Utrecht - Reactie over natuurwaarden en ecologie',
    'Dhr. De Weerd (namens Omwonenden) - Zorgen over woningbouw',
    'Scouting MERHULA - Inspraakreactie over toekomst scoutingterrein',
    'Mevr. Jonxis, Dhr. Van Hutten, Dhr. Umbgrove, Dhr. Van Motman',
    'Dhr. Van den Berg, Dhr. Van Ravels, Dhr. Asselbergs, Dhr. Buisman, Dhr. Lugtmeijer'
]
for inspreker in insprekers:
    doc.add_paragraph(f'- {inspreker}', style='List Bullet')

doc.add_heading('December 2019:', level=3)
p = doc.add_paragraph()
p.add_run('18 december 2019: Motie BOP, GroenLinks, PvdA over Participatieproces - ').bold = True
p.add_run('VERWORPEN')

# 2020
doc.add_heading('2020: Voorontwerpbestemmingsplan', level=2)

doc.add_heading('Maart-Mei 2020: Politieke onrust', level=3)
doc.add_paragraph('11 maart 2020: RIB Voortgang onderzoek sluitend krijgen businesscase')
doc.add_paragraph('20 mei 2020: RIB Vragen interpellatie Soestdijk')
doc.add_paragraph('27 mei 2020: Motie BOP over handelwijze wethouder aanbesteding')

doc.add_heading('Juli 2020: Besluitvorming', level=3)
doc.add_paragraph('30 juni - 2 juli 2020: Informatieavonden (3x)')
doc.add_paragraph('8 juli 2020: Debat')
doc.add_paragraph('15 juli 2020: VASTSTELLING VOORONTWERPBESTEMMINGSPLAN').runs[0].bold = True

doc.add_heading('Amendementen 15 juli 2020:', level=3)
doc.add_paragraph('- PvdA, GL, VVD: Wijzigingen bestemmingsplan', style='List Bullet')
doc.add_paragraph('- PvdA, GL, VVD, CU-SGP: Aanpassingen (2x)', style='List Bullet')

doc.add_heading('Moties 15 juli 2020:', level=3)
moties_2020 = [
    ('PvdA, GL, CU-SGP, VVD, D66, VoorBaarn', 'Breed gedragen'),
    ('GL, PvdA, VVD, CU-SGP', 'Ondersteuning'),
    ('GL, PvdA, D66', 'Betrekken klankbordgroep'),
    ('VVD, D66, VoorBaarn, CDA, GL', 'Coalitie + oppositie'),
    ('VVD, D66, CDA, CU-SGP, GL, PvdA', 'Brede steun')
]
for indieners, onderwerp in moties_2020:
    doc.add_paragraph(f'- {indieners}: {onderwerp}', style='List Bullet')

# 2021
doc.add_heading('2021: Ontwerpbestemmingsplan', level=2)

doc.add_heading('Maart 2021: Bezwaarschriften', level=3)
doc.add_paragraph('''De Stichting de Parel van Baarn diende bezwaar in tegen het raadsbesluit van 15 juli 2020. De onafhankelijke Bezwaarcommissie adviseerde het bezwaar niet-ontvankelijk te verklaren omdat het besluit niet vatbaar is voor bezwaar (voorbereidingsbesluit). De raad volgde dit advies op 24 maart 2021.''')

doc.add_heading('September 2021: Besluitvorming', level=3)
doc.add_paragraph('1-9 september 2021: Informatieavonden (4x)')
doc.add_paragraph('15 september 2021: Debat')
doc.add_paragraph('29 september 2021: VASTSTELLING ONTWERPBESTEMMINGSPLAN').runs[0].bold = True

doc.add_heading('Amendementen 29 september 2021:', level=3)

p = doc.add_paragraph()
p.add_run('Amendement 9I - Natuurcompensatie (D66, VVD, CDA, CU-SGP):').bold = True
doc.add_paragraph('Het college opdracht geven om in samenwerking met provincie Utrecht te zorgen voor reele en substantiele compensatie van natuur die recht doet aan de inbreuk van woningbouw in het natuurnetwerk ("het borrebos"). Compensatie moet SMART worden geformuleerd (Specifiek, Meetbaar, Acceptabel, Realistisch en Tijdgebonden).')

p = doc.add_paragraph()
p.add_run('Amendement 9J - Kwaliteitseisen restauratie (D66, VVD, CDA, CU-SGP):').bold = True
doc.add_paragraph('Het college opdracht geven om conform advies HBR Advocaten de rol van de Rijksdienst voor het Culturele Erfgoed in relatie tot monumentenvergunning en toezicht in de anterieure overeenkomst op te nemen, zodat dit ook privaatrechtelijk wordt geborgd.')

# 2022
doc.add_heading('2022: Vaststelling Bestemmingsplan', level=2)

doc.add_heading('Januari-Februari 2022: Finale besluitvorming', level=3)
doc.add_paragraph('12 januari 2022: Overeenkomst restauratie i.r.t. woningbouw Alexanderkwartier')
doc.add_paragraph('2 februari 2022: Advies HBR Advocaten, Bestemmingsplan (2x informatie)')
doc.add_paragraph('9 februari 2022: Informatie')
doc.add_paragraph('16 februari 2022: Debat')
doc.add_paragraph('23 februari 2022: BESTEMMINGSPLAN LANDGOED PALEIS SOESTDIJK DEFINITIEF VASTGESTELD').runs[0].bold = True

doc.add_heading('Amendementen 23 februari 2022:', level=3)
p = doc.add_paragraph()
p.add_run('Amendement Nokhoogte scouting (D66, VVD, VoorBaarn, CDA, GL, PvdA, CU-SGP):').bold = True
doc.add_paragraph('Nokhoogte gebouw scouting Merhula een meter hoger.')

p = doc.add_paragraph()
p.add_run('Amendement Maximaliseren bebouwing (PvdA, VoorBaarn, CDA, GL):').bold = True
doc.add_paragraph('Op de plankaart bij alle gebouwen met aanduiding (w) of (gd) in bestemming Natuur het huidige bebouwde oppervlakte als maximum opnemen.')

doc.add_heading('Moties 23 februari 2022:', level=3)
p = doc.add_paragraph()
p.add_run('Motie Onderzoek Rekenkamercommissie (PvdA, GroenLinks):').bold = True
doc.add_paragraph('Onderzoek naar inwonerparticipatie. Overwegende dat een uitgebreid participatietraject met klankbordgroepen heeft plaatsgevonden, deelnemers om uiteenlopende redenen zijn afgehaakt, veel inwoners het gevoel hadden dat inspraak geen weerklank vond bij de raad.')

p = doc.add_paragraph()
p.add_run('Motie Verduidelijken netto-opbrengsten (PvdA, GroenLinks):').bold = True
doc.add_paragraph('Verduidelijken uitgangspunten opbrengsten Alexanderkwartier. De raad heeft eerder per amendement opgedragen dat het netto resultaat van woningbouw moet worden aangewend voor renovatie van het paleis.')

doc.add_heading('Reacties externe partijen februari 2022:', level=3)
p = doc.add_paragraph()
p.add_run('Omwonenden - "Genoeg is genoeg" (4 februari 2022):').bold = True
doc.add_paragraph('Diverse stichtingen en omwonenden gaven aan NIET te gaan inspreken. Zij voelden dat inspraak geen weerklank vond bij de raad.')

p = doc.add_paragraph()
p.add_run('Bewoners Vredehofstraat/Park Vredehof/Regentesselaan (21 februari 2022):').bold = True
doc.add_paragraph('"Alstublieft geen brug. Behoed onze woonomgeving en het historische ensemble."')

doc.add_heading('Na vaststelling 2022:', level=3)
doc.add_paragraph('11 mei 2022: Schriftelijke vragen PvdA over fietsverbinding Soest-Hilversum')
doc.add_paragraph('25 mei 2022: Gedeeltelijk opheffen geheimhouding bijlagen')
doc.add_paragraph('5 oktober 2022: Verweerschrift bestemmingsplan bij Raad van State')
doc.add_paragraph('2 november 2022: COA-opvang 18 minderjarige asielzoekers op Soestdijk').runs[0].bold = True

doc.add_page_break()

# SECTIE 4: Standpunten Politieke Partijen
doc.add_heading('4. STANDPUNTEN POLITIEKE PARTIJEN', level=1)

partijen = [
    ('VVD', 'Pro-ontwikkeling, actief met amendementen voor natuurcompensatie en restauratiekwaliteit. Recent zeer kritisch op COA-opvang: vragen over incidenten, transparantie, en of de opvang moet stoppen.'),
    ('D66', 'Pro-ontwikkeling, focus op kwaliteit en natuurcompensatie. Actief met amendementen en technische vragen.'),
    ('CDA', 'Pro-ontwikkeling, maar bezorgd over financiele gevolgen uitspraak Raad van State. Vragen over gevolgen voor ambtelijke organisatie en gemeentefinancien.'),
    ('CU-SGP (ChristenUnie-SGP)', 'Pro-ontwikkeling, nadruk op kwaliteitseisen restauratie en rol Rijksdienst Cultureel Erfgoed. Vragen over vervolgproces.'),
    ('BOP (Baarnse Onafhankelijke Partij)', 'Kritisch, veel vragen over participatie, kosten, transparantie. Moties over handelwijze wethouder. Schriftelijke vragen over kosten participatietraject.'),
    ('GroenLinks', 'Focus op natuur, participatie, kritisch op proces. Moties samen met PvdA voor rekenkameronderzoek. Vragen over vervolgproces en omgevingsvisie.'),
    ('PvdA', 'Focus op participatie, inwonersinspraak, rekenkameronderzoek. Veel technische vragen over bestemmingsplan. Moties voor verduidelijking opbrengsten.'),
    ('VoorBaarn', 'Constructief-kritisch, actief met amendementen samen met andere partijen.'),
    ('Lijst Schouten', 'Recent actief (2024) samen met VVD.')
]

for partij, standpunt in partijen:
    p = doc.add_paragraph()
    p.add_run(f'{partij}: ').bold = True
    p.add_run(standpunt)
    doc.add_paragraph()

doc.add_page_break()

# SECTIE 5: Standpunten Externe Organisaties
doc.add_heading('5. STANDPUNTEN EXTERNE ORGANISATIES', level=1)

doc.add_heading('Stichting de Parel van Baarn', level=2)
doc.add_paragraph('''Volledige naam: Stichting tot behoud van het historisch ensemble Paleis Soestdijk "De Parel van Baarn"
Voorzitter: Mr. M.L.M. van Ravels

Standpunten:
- Kritisch op afwijking van oorspronkelijke plannen door MeyerBergman
- Tegen bouwen in beschermde natuur (ook klein deel Borrebos)
- Tegen hotel op de parade
- Zorgen over financiele gegoedheid MeyerBergman
- Pleit voor rechtszekerheid en betrouwbare overheid

Inspraak 8 september 2021 (citaat):
"Ligt die papierbrij niet grotendeels aan Meijer Bergman, die van zijn oorspronkelijke plan is afgeweken? Het is goed dat inmiddels het aantal flats in het Borrebos - onder publieke druk - is teruggeschroefd. Maar bouwen in de beschermde natuur - ook een klein deel daarvan - blijft voor ons onbespreekbaar."

Brief maart 2025:
Zorgen over verkoop Intendance (Parade/Herencluster) en of opbrengst daadwerkelijk restauratie ten goede komt. Pleit voor storting op escrow-account als garantie.''')

doc.add_heading('Natuur en Milieufederatie Utrecht', level=2)
doc.add_paragraph('''Standpunten:
- Focus op natuurwaarden en ecologie
- Zorgen over woningbouw in Natuurnetwerk Nederland
- Pleit voor adequate natuurcompensatie
- Actief met inspraakreacties bij alle beslismomenten''')

doc.add_heading('Omwonenden', level=2)
doc.add_paragraph('''Diverse omwonenden hebben ingesproken, waaronder:
- Bewoners Vredehofstraat/Park Vredehof/Regentesselaan (Soest)
- Diverse individuele insprekers

Standpunten:
- Zorgen over woningbouw en verdichting
- Tegen aanleg brug
- "Genoeg is genoeg" - gevoel dat inspraak geen weerklank vindt
- Behoud historisch ensemble en woonomgeving''')

doc.add_heading('Scouting MERHULA', level=2)
doc.add_paragraph('''Standpunten:
- Zorgen over toekomst scoutingterrein
- Actief met inspraakreacties
- Amendement 2022 verhoogde nokhoogte scoutinggebouw met 1 meter''')

doc.add_heading('Stichting Behoud het Borrebos', level=2)
doc.add_paragraph('''Standpunten:
- Kritisch op COA-gebruik marechausseeterrein
- Stelt dat omgevingsvergunning ontbreekt
- Heeft gemeente in gebreke gesteld
- Verzoek om handhaving of beeindigen oneigenlijk gebruik''')

doc.add_heading('HBR Advocaten', level=2)
doc.add_paragraph('''Rol: Onafhankelijk juridisch adviseur ingehuurd door gemeente
Adviezen:
- Beoordeling anterieure overeenkomst
- Rol Rijksdienst Cultureel Erfgoed
- Juridische borging restauratiekwaliteit
- Reactie op bevindingen door college''')

doc.add_heading('Staatsbosbeheer', level=2)
doc.add_paragraph('''Rol: Betrokken bij grondruil
Relevante documenten:
- Concept overeenkomst van ruiling MBE en Staatsbosbeheer (juni 2021)
- Brief mbt grondruil Soestdijk - Didam-arrest (januari 2022)''')

doc.add_heading('MeyerBergman Erfgoed Groep (MBEG/MBE)', level=2)
doc.add_paragraph('''Rol: Eigenaar en ontwikkelaar sinds december 2017

Made by Holland plan:
- Restauratie paleis, park en bos
- Woningbouw Alexanderkwartier ter financiering
- Publieke toegankelijkheid
- Culturele en zakelijke evenementen

Actueel:
- Start restauratie buitenkant paleis: januari/februari 2026
- Voorbereiding restauratie binnenkant (monumentenvergunning nodig)
- Presentatie geactualiseerde visie: april 2025''')

doc.add_page_break()

# SECTIE 6: COA-opvang
doc.add_heading('6. COA-OPVANG OP HET MARECHAUSSEETERREIN', level=1)

doc.add_heading('Achtergrond', level=2)
doc.add_paragraph('''Het voormalige Marechausseeterrein bij Paleis Soestdijk was oorspronkelijk bestemd voor woningbouw (Alexanderkwartier) om de restauratie van het paleis te bekostigen. Na de vernietiging van het bestemmingsplan door de Raad van State is dit terrein tijdelijk in gebruik genomen door het COA.''')

doc.add_heading('Chronologie COA-opvang', level=2)

coa_tijdlijn = [
    ('Juni 2022', 'Start opvang op marechausseeterrein'),
    ('Juli 2022', '50 alleenstaande minderjarige vluchtelingen (AMVers) opgevangen'),
    ('20 september 2022', 'COA verzoekt uitbreiding met 18 plekken'),
    ('5 oktober 2022', 'College informeert raad over uitbreiding'),
    ('November 2022', 'Uitbreiding met 18 minderjarige asielzoekers goedgekeurd'),
    ('Februari 2025', 'Inmiddels 152 asielzoekers opgevangen'),
    ('2025', 'College start procedure voor definitieve vestiging COA'),
    ('November 2025', 'Vragen over incidenten en aanhoudingen'),
    ('Januari 2026', 'Discussie over transparantie huuropbrengsten')
]

for datum, gebeurtenis in coa_tijdlijn:
    p = doc.add_paragraph()
    p.add_run(f'{datum}: ').bold = True
    p.add_run(gebeurtenis)

doc.add_heading('RIB 5 oktober 2022 - Uitbreiding opvang', level=2)
doc.add_paragraph('''Portefeuillehouder: Wethouder De Vries

Inhoud:
"Sinds medio juli worden 50 alleenstaande minderjarige vluchtelingen (AMVers) opgevangen op het terrein van de voormalige marechausseekazerne bij Paleis Soestdijk. Het COA heeft op 20 september jl. het formele verzoek gedaan aan de gemeente dit aantal uit te breiden met 18 opvangplekken, gelet op de huidige situatie in Ter Apel.

De centrale locatie in Ter Apel waar asielzoekers worden opgevangen wordt niet beschouwd als een geschikte plek voor alleenstaande minderjarigen. Het COA zoekt daarom continu naar nieuwe opvangmogelijkheden en het efficienter gebruiken van bestaande opvanglocaties. Baarn wordt door het COA gezien als een locatie die efficienter gebruikt kan worden.

Het COA geeft de garantie dat er ook voor deze extra groep jongeren voldoende begeleiding wordt gegeven. In het hoofdgebouw zijn klaslokalen ingericht en wordt les gegeven door Het Element (Taalcentrum in Amersfoort)."''')

doc.add_heading('Mondelinge vragen VVD - 19 februari 2025', level=2)
doc.add_paragraph('''"Sinds juni 2022 wordt het voormalige Marechausseeterrein door het COA gebruikt als asielzoekerscentrum. Aanvankelijk lag het in de bedoeling dat de opvang voor de duur van maximaal een jaar zou zijn. We zijn bijna drie jaar verder en er worden inmiddels 152 asielzoekers opgevangen.

Aanvullend heeft het College aangegeven de procedure op te starten om op het Marechausseeterrein definitief een COA te willen vestigen voor onbepaalde tijd.

Via de media begrepen wij dat Stichting Behoud het Borrebos van oordeel is dat het nooit tot een afwijkingsvergunning is gekomen. De stichting heeft de gemeente schriftelijk in gebreke gesteld."''')

doc.add_heading('Mondelinge vragen VVD - 26 november 2025', level=2)
doc.add_paragraph('''Vragen over de toezegging van de wethouder:
a. Acht de wethouder dat het COA er "een potje van maakt", gelet op de recente incidenten en herhaalde aanhoudingen?
b. Welke stappen zijn inmiddels in gang gezet om daadwerkelijk te stoppen met de huidige opvang?
c. Welke objectieve maatstaven hanteert het college om te bepalen wanneer de grens is bereikt?

Vragen over informatievoorziening:
- Waarom is het presidium wel geinformeerd en de raad niet?
- Erkent het college dat hiermee de indruk wordt gewekt dat informatie wordt achtergehouden?

Vragen over veiligheid en samenstelling:
- Is het college volledig geinformeerd over de vechtpartijen tussen groepen bewoners?
- Hoeveel bewoners verblijven momenteel? Hoeveel minderjarigen, gezinnen, alleenstaanden?
- Hoeveel minderjarige meisjes verblijven er en welke maatregelen zijn genomen voor hun veiligheid?''')

doc.add_page_break()

# SECTIE 7: Raad van State
doc.add_heading('7. RAAD VAN STATE UITSPRAAK EN GEVOLGEN', level=1)

doc.add_heading('Uitspraak januari 2024', level=2)
doc.add_paragraph('''De Raad van State heeft in januari 2024 delen van het bestemmingsplan Landgoed Paleis Soestdijk vernietigd. Dit heeft grote gevolgen voor de geplande ontwikkelingen.''')

doc.add_heading('Mondelinge vragen CDA - 31 januari 2024', level=2)
doc.add_paragraph('''"De uitspraak van de Raad van State is voor ons reden voor grote zorg. Het achteloze van de hand doen destijds van dit (koninklijk) cultureel erfgoed door de Rijksoverheid heeft Baarn in materieel en immaterieel opzicht onevenredig veel gekost.

Het is evident dat er opnieuw veel gevraagd gaat worden van ons ambtenarenapparaat, van ons college, van de gemeenteraad, dat belanghebbenden en deskundigen moeten worden geraadpleegd, kortom: dat dit opnieuw veel capaciteit en veel geld zal gaan kosten."''')

doc.add_heading('RIB Gevolgen uitspraak - 29 februari 2024', level=2)
doc.add_paragraph('''"De vernietiging betekent vertraging van de restauratie van paleis, park en bos. De afgelopen jaren is het landgoed Paleis Soestdijk technisch in stand gehouden. De MeyerBergman Erfgoed Groep (MBEG) heeft aangegeven dat het noodzakelijke onderhoud ook de komende periode zal worden gecontinueerd."''')

doc.add_heading('Reflectiedocument - 25 september 2024', level=2)
doc.add_paragraph('''De gemeenteraad heeft in juni 2024 een reflectiebijeenkomst gehouden over de eigen rol rondom de besluitvorming.

Doel: Reflecteren op eigen ervaringen en beelden over het besluitvormingsproces delen.

De bijeenkomst was op een externe locatie met externe begeleider (Koos Janssen, oud-burgemeester Zeist).

Besluit raad 25 september 2024:
1. Het reflectiedocument vaststellen
2. Het presidium verzoeken eventuele vervolgacties voor te bereiden''')

doc.add_page_break()

# SECTIE 8: Huidige Stand van Zaken
doc.add_heading('8. HUIDIGE STAND VAN ZAKEN (2025-2026)', level=1)

doc.add_heading('RIB Stand van zaken - 18 december 2025', level=2)
doc.add_paragraph('''Restauratie:
- De Naald aan de Torenlaan wordt gerestaureerd (wordt afgerond)
- Start restauratie buitenkant paleis: januari/februari 2026
- Fasegewijs werken zodat activiteiten kunnen doorgaan
- Geen omgevingsvergunning nodig voor buitenkant
- Voorbereiding restauratie binnenkant (monumentenvergunning nodig)''')

doc.add_heading('Peiling VVD - 21 januari 2026', level=2)
doc.add_paragraph('''Peilpunt 1 - Transparantie:
"Deelt de raad de mening dat inzicht nodig is in de hoogte van de huuropbrengsten uit verhuur aan het COA, zodat kan worden geborgd dat deze worden aangewend voor renovatie van Paleis Soestdijk?"

Peilpunt 2 - Participatie:
"Deelt de raad de mening dat inwoners van Baarn en Soest actief worden geinformeerd en in de gelegenheid worden gesteld hun zienswijze te geven over de plannen en de opvang?"

Aanleiding: "De wethouder heeft aangegeven niet te weten - en ook nadrukkelijk niet te willen weten - wat de hoogte van deze huuropbrengsten is."''')

doc.add_page_break()

# SECTIE 9: Bronnen
doc.add_heading('9. BRONNEN EN DOCUMENTEN', level=1)

doc.add_heading('Belangrijkste documenten', level=2)
bronnen = [
    'Omgevingsvisie Paleis Soestdijk 2012',
    'Ronde Tafel-advies 2015',
    'Made by Holland Plan voor herontwikkeling (2015-2016)',
    'Koopovereenkomst 3 juli 2017',
    'Akte van levering 20 december 2017',
    'Ruimtelijk Kader Landgoed Paleis Soestdijk (april 2019)',
    'Voorontwerpbestemmingsplan Landgoed Paleis Soestdijk (juli 2020)',
    'Adviezen HBR Advocaten (2021-2022)',
    'Ontwerpbestemmingsplan Landgoed Paleis Soestdijk (september 2021)',
    'Bestemmingsplan Landgoed Paleis Soestdijk - Vastgesteld (23 februari 2022)',
    'Verweerschrift bestemmingsplan bij Raad van State (oktober 2022)',
    'RIB Uitbreiding opvang COA Soestdijk (oktober 2022)',
    'Uitspraak Raad van State (januari 2024)',
    'RIB Gevolgen uitspraak Raad van State (februari 2024)',
    'Reflectiedocument besluitvormingsproces (september 2024)',
    'RIB Stand van zaken ontwikkelingen Paleis Soestdijk (december 2025)'
]
for bron in bronnen:
    doc.add_paragraph(f'- {bron}', style='List Bullet')

doc.add_heading('Database informatie', level=2)
doc.add_paragraph('''Dit document is samengesteld op basis van de Baarn Raadsinformatie database:
- 659 vergaderingen
- 7.068 agenda items
- 16.763 documenten
- Periode: december 2009 - januari 2026

Specifiek voor Paleis Soestdijk:
- 44 agenda items
- 295 documenten''')

doc.add_paragraph()
p = doc.add_paragraph('Document gegenereerd: januari 2026')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph('Bron: Baarn Raadsinformatie MCP Server')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Opslaan
output_path = 'data/Paleis_Soestdijk_Volledige_Tijdlijn_en_COA.docx'
doc.save(output_path)
print(f'Document opgeslagen: {output_path}')
