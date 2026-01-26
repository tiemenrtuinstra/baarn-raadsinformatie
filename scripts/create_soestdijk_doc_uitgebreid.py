# -*- coding: utf-8 -*-
import json
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Laad de documenten
with open('data/soestdijk_docs.json', 'r', encoding='utf-8') as f:
    docs = json.load(f)

doc = Document()

# Stel standaard lettertype in
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)

# Titel
title = doc.add_heading('PALEIS SOESTDIJK', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Volledige Documentatie: Moties, Amendementen, Inspraakreacties en Standpunten')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.runs[0]
subtitle_run.font.size = Pt(14)
subtitle_run.font.italic = True

doc.add_paragraph('Gemeente Baarn - Raadsinformatie 2011-2026')
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Inhoudsopgave
doc.add_heading('INHOUDSOPGAVE', level=1)
toc_items = [
    '1. Overzicht en Statistieken',
    '2. Amendementen (volledige teksten)',
    '3. Moties (volledige teksten)',
    '4. Inspraakreacties en Zienswijzen',
    '5. Brieven en Reacties Externe Partijen',
    '6. Juridische Adviezen (HBR Advocaten)',
    '7. Alle Insprekers per Vergadering'
]
for item in toc_items:
    doc.add_paragraph(item)

doc.add_page_break()

# SECTIE 1: Overzicht
doc.add_heading('1. OVERZICHT EN STATISTIEKEN', level=1)

doc.add_paragraph(f'''Dit document bevat de volledige teksten van {len(docs)} documenten over Paleis Soestdijk, inclusief:
- Amendementen met volledige tekst en toelichting
- Moties met overwegingen en dictum
- Inspraakreacties van burgers en organisaties
- Brieven van externe partijen
- Juridische adviezen

Bronnen: Baarn Raadsinformatie Database (2009-2026)
- 44 agenda items over Paleis Soestdijk
- 295 documenten over Paleis Soestdijk
- 60 inspraakreacties
- 13 moties
- 15 amendementen''')

doc.add_page_break()

# SECTIE 2: Amendementen
doc.add_heading('2. AMENDEMENTEN (VOLLEDIGE TEKSTEN)', level=1)

doc.add_paragraph('Hieronder volgen alle amendementen over Paleis Soestdijk met volledige tekst.')
doc.add_paragraph()

# Filter amendementen
amendementen = [d for d in docs if 'amendement' in d['title'].lower()]

for am in amendementen:
    # Skip bijlagen die geen echte amendementen zijn
    if 'bijlage' in am['title'].lower() and 'aangepast' in am['title'].lower():
        continue

    doc.add_heading(f"Amendement: {am['date']}", level=2)
    p = doc.add_paragraph()
    p.add_run(am['title']).bold = True
    doc.add_paragraph()

    # Voeg inhoud toe (beperk voor leesbaarheid)
    content = am['content'][:8000] if am['content'] else 'Geen inhoud beschikbaar'
    # Vervang problematische karakters
    content = content.replace('\x00', '').replace('\r', '')
    doc.add_paragraph(content)
    doc.add_paragraph()
    doc.add_paragraph('_' * 80)
    doc.add_paragraph()

doc.add_page_break()

# SECTIE 3: Moties
doc.add_heading('3. MOTIES (VOLLEDIGE TEKSTEN)', level=1)

doc.add_paragraph('Hieronder volgen alle moties over Paleis Soestdijk met volledige tekst.')
doc.add_paragraph()

# Filter moties
moties = [d for d in docs if 'motie' in d['title'].lower()]

for mo in moties:
    # Skip bijlagen
    if 'bijlage' in mo['title'].lower():
        continue

    doc.add_heading(f"Motie: {mo['date']}", level=2)
    p = doc.add_paragraph()
    p.add_run(mo['title']).bold = True
    doc.add_paragraph()

    content = mo['content'][:8000] if mo['content'] else 'Geen inhoud beschikbaar'
    content = content.replace('\x00', '').replace('\r', '')
    doc.add_paragraph(content)
    doc.add_paragraph()
    doc.add_paragraph('_' * 80)
    doc.add_paragraph()

doc.add_page_break()

# SECTIE 4: Inspraakreacties
doc.add_heading('4. INSPRAAKREACTIES EN ZIENSWIJZEN', level=1)

doc.add_paragraph('Hieronder volgen inspraakreacties van burgers en organisaties.')
doc.add_paragraph()

# Filter inspraakreacties
inspreek = [d for d in docs if 'inspraa' in d['title'].lower() or 'inspreek' in d['title'].lower() or 'zienswijze' in d['title'].lower()]

for ins in inspreek:
    doc.add_heading(f"Inspraak: {ins['date']}", level=2)
    p = doc.add_paragraph()
    p.add_run(ins['title']).bold = True
    doc.add_paragraph()

    content = ins['content'][:8000] if ins['content'] else 'Geen inhoud beschikbaar'
    content = content.replace('\x00', '').replace('\r', '')
    doc.add_paragraph(content)
    doc.add_paragraph()
    doc.add_paragraph('_' * 80)
    doc.add_paragraph()

doc.add_page_break()

# SECTIE 5: Brieven externe partijen
doc.add_heading('5. BRIEVEN EN REACTIES EXTERNE PARTIJEN', level=1)

doc.add_paragraph('Hieronder volgen brieven en reacties van externe partijen.')
doc.add_paragraph()

# Filter brieven en reacties
brieven = [d for d in docs if ('brief' in d['title'].lower() or 'reactie' in d['title'].lower() or 'parel' in d['title'].lower() or 'omwonenden' in d['title'].lower())
           and 'inspraa' not in d['title'].lower()]

for br in brieven:
    doc.add_heading(f"Brief/Reactie: {br['date']}", level=2)
    p = doc.add_paragraph()
    p.add_run(br['title']).bold = True
    doc.add_paragraph()

    content = br['content'][:10000] if br['content'] else 'Geen inhoud beschikbaar'
    content = content.replace('\x00', '').replace('\r', '')
    doc.add_paragraph(content)
    doc.add_paragraph()
    doc.add_paragraph('_' * 80)
    doc.add_paragraph()

doc.add_page_break()

# SECTIE 6: Juridische adviezen
doc.add_heading('6. JURIDISCHE ADVIEZEN (HBR ADVOCATEN)', level=1)

doc.add_paragraph('Hieronder volgen de juridische adviezen van HBR Advocaten.')
doc.add_paragraph()

# Filter HBR documenten
hbr = [d for d in docs if 'hbr' in d['title'].lower() or 'advocat' in d['title'].lower()]

for h in hbr:
    doc.add_heading(f"Juridisch Advies: {h['date']}", level=2)
    p = doc.add_paragraph()
    p.add_run(h['title']).bold = True
    doc.add_paragraph()

    content = h['content'][:10000] if h['content'] else 'Geen inhoud beschikbaar'
    content = content.replace('\x00', '').replace('\r', '')
    doc.add_paragraph(content)
    doc.add_paragraph()
    doc.add_paragraph('_' * 80)
    doc.add_paragraph()

doc.add_page_break()

# SECTIE 7: Alle insprekers
doc.add_heading('7. OVERZICHT ALLE INSPREKERS', level=1)

doc.add_heading('Ruimtelijk Kader - April 2019', level=2)
insprekers_2019 = [
    'Stichting de Parel van Baarn - Open brief behoud historisch ensemble',
    'Natuur en Milieufederatie Utrecht - Reactie natuurwaarden',
    'Dhr. De Weerd en Storm (namens Omwonenden)',
    'Mevr. Jonxis',
    'Dhr. Van Hutten',
    'Scouting MERHULA (2 inspraakreacties)',
    'Dhr. Umbgrove',
    'Dhr. Van Motman (VVP)',
    'Dhr. Van den Berg (CCLV)',
    'Dhr. Van Ravels (Parel van Baarn)',
    'Dhr. Asselbergs',
    'Dhr. Buisman',
    'Dhr. Lugtmeijer (Stichting)'
]
for ins in insprekers_2019:
    doc.add_paragraph(f'- {ins}', style='List Bullet')

doc.add_heading('Voorontwerpbestemmingsplan - Juli 2020', level=2)
insprekers_2020 = [
    'Dhr. Waagepetersen',
    'Dhr. B. Smit',
    'Mevr. Coumont',
    'Dhr. Van Assema',
    'Dhr. De Weerd',
    'Dhr. M. Smit',
    'Mevr. Geerts',
    'Mevr. Jonxis',
    'Dhr. Luth',
    'Dhr. Van Hutten',
    'Dhr. Koolma',
    'Dhr. Van Motman',
    'Dhr. Ten Broeke',
    'Mevr. Broodbakker',
    'Dhr. Van Ommeren',
    'Dhr. Wiltink',
    'Dhr. Asselbergs',
    'Mevr. Beekmans'
]
for ins in insprekers_2020:
    doc.add_paragraph(f'- {ins}', style='List Bullet')

doc.add_heading('Bestemmingsplan - Februari 2022', level=2)
insprekers_2022 = [
    'Scouting Merhula',
    'Inwoner (geanonimiseerd)',
    'Mevr. De Vrey-Vringer',
    'Omwonenden collectief - "Genoeg is genoeg"',
    'Bewoners Vredehofstraat/Park Vredehof/Regentesselaan',
    'MeyerBergman Erfgoed Groep (MBE)'
]
for ins in insprekers_2022:
    doc.add_paragraph(f'- {ins}', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph('Document gegenereerd: januari 2026')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph('Bron: Baarn Raadsinformatie MCP Server')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Opslaan
output_path = 'data/Paleis_Soestdijk_UITGEBREID_met_volledige_teksten.docx'
doc.save(output_path)
print(f'Document opgeslagen: {output_path}')
