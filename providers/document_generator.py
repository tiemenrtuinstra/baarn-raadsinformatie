#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Document Generator voor Baarn Politiek MCP Server.
Genereert Word documenten voor moties en amendementen.
"""

from pathlib import Path
from datetime import date
from typing import Dict, List, Optional
import re

from core.config import Config
from shared.logging_config import get_logger

logger = get_logger('document-generator')

# Probeer python-docx te importeren
try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning('python-docx not installed. Document generation will be limited to markdown.')


class DocumentGenerator:
    """Generator voor officiële gemeente documenten (moties en amendementen)."""

    def __init__(self):
        self.output_dir = Config.DATA_DIR / 'generated'
        self.output_dir.mkdir(parents=True, exist_ok=True)
        logger.info(f'DocumentGenerator initialized. Output dir: {self.output_dir}')

    def _sanitize_filename(self, text: str, max_length: int = 40) -> str:
        """Maak een veilige bestandsnaam van tekst."""
        # Verwijder ongeldige karakters
        safe = re.sub(r'[<>:"/\\|?*]', '', text)
        # Vervang spaties door underscores
        safe = safe.replace(' ', '_')
        # Beperk lengte
        return safe[:max_length]

    def _generate_filename(self, doc_type: str, titel: str) -> str:
        """Genereer bestandsnaam voor document."""
        safe_title = self._sanitize_filename(titel)
        timestamp = date.today().isoformat()
        return f"{doc_type}_{timestamp}_{safe_title}.docx"

    def generate_motie(
        self,
        titel: str,
        indieners: List[str],
        partijen: List[str],
        constateringen: List[str],
        overwegingen: List[str],
        verzoeken: List[str],
        vergadering_datum: str = None,
        agendapunt: str = None,
        toelichting: str = None
    ) -> Dict:
        """
        Genereer een motie document in Notubiz-conform formaat.

        Args:
            titel: Titel van de motie
            indieners: Lijst van indieners (namen)
            partijen: Lijst van partijen
            constateringen: "Constaterende dat" punten
            overwegingen: "Overwegende dat" punten
            verzoeken: "Verzoekt het college" punten
            vergadering_datum: Datum vergadering (optioneel)
            agendapunt: Agendapunt nummer (optioneel)
            toelichting: Optionele toelichting

        Returns:
            Dict met filepath en markdown content
        """
        logger.info(f'Generating motie: {titel}')

        # Genereer markdown versie (altijd beschikbaar)
        markdown = self._generate_motie_markdown(
            titel, indieners, partijen, constateringen,
            overwegingen, verzoeken, vergadering_datum,
            agendapunt, toelichting
        )

        result = {
            'titel': titel,
            'type': 'motie',
            'markdown': markdown
        }

        # Genereer Word document als beschikbaar
        if DOCX_AVAILABLE:
            filepath = self._generate_motie_docx(
                titel, indieners, partijen, constateringen,
                overwegingen, verzoeken, vergadering_datum,
                agendapunt, toelichting
            )
            result['filepath'] = str(filepath)
            result['filename'] = filepath.name
            logger.info(f'Motie generated: {filepath}')
        else:
            # Sla markdown op als fallback
            filename = self._generate_filename('motie', titel).replace('.docx', '.md')
            filepath = self.output_dir / filename
            filepath.write_text(markdown, encoding='utf-8')
            result['filepath'] = str(filepath)
            result['filename'] = filepath.name
            result['warning'] = 'python-docx not installed, generated markdown instead'

        return result

    def _generate_motie_markdown(
        self,
        titel: str,
        indieners: List[str],
        partijen: List[str],
        constateringen: List[str],
        overwegingen: List[str],
        verzoeken: List[str],
        vergadering_datum: str = None,
        agendapunt: str = None,
        toelichting: str = None
    ) -> str:
        """Genereer markdown versie van motie."""
        lines = []
        lines.append('# MOTIE')
        lines.append('')

        if vergadering_datum:
            lines.append(f'**Vergadering:** {vergadering_datum}')
        if agendapunt:
            lines.append(f'**Agendapunt:** {agendapunt}')
        lines.append('')

        lines.append(f'## Motie: {titel}')
        lines.append('')

        lines.append(f'**Ingediend door:** {", ".join(indieners)}')
        lines.append(f'**Namens:** {", ".join(partijen)}')
        lines.append('')

        lines.append('---')
        lines.append('')
        lines.append('*De raad van de gemeente Baarn, in vergadering bijeen,*')
        lines.append('')

        lines.append('### Constaterende dat:')
        for item in constateringen:
            lines.append(f'- {item}')
        lines.append('')

        lines.append('### Overwegende dat:')
        for item in overwegingen:
            lines.append(f'- {item}')
        lines.append('')

        lines.append('### Verzoekt het college:')
        for item in verzoeken:
            lines.append(f'- {item}')
        lines.append('')

        if toelichting:
            lines.append('### Toelichting')
            lines.append(toelichting)
            lines.append('')

        lines.append('*en gaat over tot de orde van de dag.*')
        lines.append('')
        lines.append('---')
        lines.append('')

        # Ondertekening
        lines.append('### Ondertekening')
        lines.append('')
        for indiener, partij in zip(indieners, partijen):
            lines.append(f'________________________')
            lines.append(f'{indiener} ({partij})')
            lines.append('')

        return '\n'.join(lines)

    def _generate_motie_docx(
        self,
        titel: str,
        indieners: List[str],
        partijen: List[str],
        constateringen: List[str],
        overwegingen: List[str],
        verzoeken: List[str],
        vergadering_datum: str = None,
        agendapunt: str = None,
        toelichting: str = None
    ) -> Path:
        """Genereer Word document voor motie."""
        doc = Document()

        # Stijlen instellen
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        # Header rechts uitgelijnd
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header.add_run('Gemeente Baarn')
        run.bold = True

        if vergadering_datum:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run(f'Vergadering: {vergadering_datum}')

        if agendapunt:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run(f'Agendapunt: {agendapunt}')

        doc.add_paragraph()

        # Titel
        title_para = doc.add_paragraph()
        title_run = title_para.add_run('MOTIE')
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        subtitle = doc.add_paragraph()
        sub_run = subtitle.add_run(f'Motie: {titel}')
        sub_run.bold = True
        sub_run.font.size = Pt(14)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Indieners
        doc.add_paragraph(f'Ingediend door: {", ".join(indieners)}')
        doc.add_paragraph(f'Namens: {", ".join(partijen)}')

        doc.add_paragraph()

        # Intro
        intro = doc.add_paragraph()
        intro_run = intro.add_run('De raad van de gemeente Baarn, in vergadering bijeen,')
        intro_run.italic = True

        doc.add_paragraph()

        # Constaterende dat
        const_title = doc.add_paragraph()
        const_title.add_run('Constaterende dat:').bold = True
        for item in constateringen:
            bullet = doc.add_paragraph(style='List Bullet')
            bullet.add_run(item)

        doc.add_paragraph()

        # Overwegende dat
        overw_title = doc.add_paragraph()
        overw_title.add_run('Overwegende dat:').bold = True
        for item in overwegingen:
            bullet = doc.add_paragraph(style='List Bullet')
            bullet.add_run(item)

        doc.add_paragraph()

        # Verzoekt het college
        verz_title = doc.add_paragraph()
        verz_title.add_run('Verzoekt het college:').bold = True
        for item in verzoeken:
            bullet = doc.add_paragraph(style='List Bullet')
            bullet.add_run(item)

        # Toelichting
        if toelichting:
            doc.add_paragraph()
            toel_title = doc.add_paragraph()
            toel_title.add_run('Toelichting:').bold = True
            doc.add_paragraph(toelichting)

        doc.add_paragraph()

        # Sluiting
        closing = doc.add_paragraph()
        closing_run = closing.add_run('en gaat over tot de orde van de dag.')
        closing_run.italic = True

        doc.add_paragraph()
        doc.add_paragraph()

        # Ondertekening
        for indiener, partij in zip(indieners, partijen):
            doc.add_paragraph('_' * 40)
            doc.add_paragraph(f'{indiener} ({partij})')
            doc.add_paragraph()

        # Opslaan
        filename = self._generate_filename('motie', titel)
        filepath = self.output_dir / filename
        doc.save(filepath)

        return filepath

    def generate_amendement(
        self,
        titel: str,
        indieners: List[str],
        partijen: List[str],
        raadsvoorstel_nummer: str,
        raadsvoorstel_titel: str,
        wijzigingen: List[Dict[str, str]],
        toelichting: str = None,
        vergadering_datum: str = None,
        agendapunt: str = None
    ) -> Dict:
        """
        Genereer een amendement document in Notubiz-conform formaat.

        Args:
            titel: Titel van het amendement
            indieners: Lijst van indieners
            partijen: Lijst van partijen
            raadsvoorstel_nummer: Nummer van het raadsvoorstel
            raadsvoorstel_titel: Titel van het raadsvoorstel
            wijzigingen: Lijst van tekstwijzigingen [{'oorspronkelijk': '...', 'wordt': '...'}]
            toelichting: Optionele toelichting
            vergadering_datum: Datum vergadering
            agendapunt: Agendapunt nummer

        Returns:
            Dict met filepath en markdown content
        """
        logger.info(f'Generating amendement: {titel}')

        # Genereer markdown versie
        markdown = self._generate_amendement_markdown(
            titel, indieners, partijen, raadsvoorstel_nummer,
            raadsvoorstel_titel, wijzigingen, toelichting,
            vergadering_datum, agendapunt
        )

        result = {
            'titel': titel,
            'type': 'amendement',
            'markdown': markdown
        }

        # Genereer Word document als beschikbaar
        if DOCX_AVAILABLE:
            filepath = self._generate_amendement_docx(
                titel, indieners, partijen, raadsvoorstel_nummer,
                raadsvoorstel_titel, wijzigingen, toelichting,
                vergadering_datum, agendapunt
            )
            result['filepath'] = str(filepath)
            result['filename'] = filepath.name
            logger.info(f'Amendement generated: {filepath}')
        else:
            filename = self._generate_filename('amendement', titel).replace('.docx', '.md')
            filepath = self.output_dir / filename
            filepath.write_text(markdown, encoding='utf-8')
            result['filepath'] = str(filepath)
            result['filename'] = filepath.name
            result['warning'] = 'python-docx not installed, generated markdown instead'

        return result

    def _generate_amendement_markdown(
        self,
        titel: str,
        indieners: List[str],
        partijen: List[str],
        raadsvoorstel_nummer: str,
        raadsvoorstel_titel: str,
        wijzigingen: List[Dict[str, str]],
        toelichting: str = None,
        vergadering_datum: str = None,
        agendapunt: str = None
    ) -> str:
        """Genereer markdown versie van amendement."""
        lines = []
        lines.append('# AMENDEMENT')
        lines.append('')

        if vergadering_datum:
            lines.append(f'**Vergadering:** {vergadering_datum}')
        if agendapunt:
            lines.append(f'**Agendapunt:** {agendapunt}')
        lines.append('')

        lines.append(f'## Amendement: {titel}')
        lines.append('')

        lines.append(f'**Ingediend door:** {", ".join(indieners)}')
        lines.append(f'**Namens:** {", ".join(partijen)}')
        lines.append('')

        lines.append(f'**Betreft raadsvoorstel:** {raadsvoorstel_nummer} - {raadsvoorstel_titel}')
        lines.append('')

        lines.append('---')
        lines.append('')
        lines.append('*De raad van de gemeente Baarn, in vergadering bijeen,*')
        lines.append('')

        lines.append('### Besluit:')
        lines.append('')
        lines.append('Het raadsvoorstel als volgt te wijzigen:')
        lines.append('')

        for i, wijziging in enumerate(wijzigingen, 1):
            lines.append(f'#### Wijziging {i}')
            lines.append('')
            lines.append('**De tekst:**')
            lines.append(f'> "{wijziging.get("oorspronkelijk", "")}"')
            lines.append('')
            lines.append('**Te wijzigen in:**')
            lines.append(f'> "{wijziging.get("wordt", "")}"')
            lines.append('')

        if toelichting:
            lines.append('### Toelichting')
            lines.append(toelichting)
            lines.append('')

        lines.append('---')
        lines.append('')

        # Ondertekening
        lines.append('### Ondertekening')
        lines.append('')
        for indiener, partij in zip(indieners, partijen):
            lines.append(f'________________________')
            lines.append(f'{indiener} ({partij})')
            lines.append('')

        return '\n'.join(lines)

    def _generate_amendement_docx(
        self,
        titel: str,
        indieners: List[str],
        partijen: List[str],
        raadsvoorstel_nummer: str,
        raadsvoorstel_titel: str,
        wijzigingen: List[Dict[str, str]],
        toelichting: str = None,
        vergadering_datum: str = None,
        agendapunt: str = None
    ) -> Path:
        """Genereer Word document voor amendement."""
        doc = Document()

        # Stijlen instellen
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        # Header
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header.add_run('Gemeente Baarn')
        run.bold = True

        if vergadering_datum:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run(f'Vergadering: {vergadering_datum}')

        if agendapunt:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run(f'Agendapunt: {agendapunt}')

        doc.add_paragraph()

        # Titel
        title_para = doc.add_paragraph()
        title_run = title_para.add_run('AMENDEMENT')
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        subtitle = doc.add_paragraph()
        sub_run = subtitle.add_run(f'Amendement: {titel}')
        sub_run.bold = True
        sub_run.font.size = Pt(14)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Indieners
        doc.add_paragraph(f'Ingediend door: {", ".join(indieners)}')
        doc.add_paragraph(f'Namens: {", ".join(partijen)}')

        doc.add_paragraph()

        # Raadsvoorstel referentie
        ref = doc.add_paragraph()
        ref.add_run('Betreft raadsvoorstel: ').bold = True
        ref.add_run(f'{raadsvoorstel_nummer} - {raadsvoorstel_titel}')

        doc.add_paragraph()

        # Intro
        intro = doc.add_paragraph()
        intro_run = intro.add_run('De raad van de gemeente Baarn, in vergadering bijeen,')
        intro_run.italic = True

        doc.add_paragraph()

        # Besluit
        besluit = doc.add_paragraph()
        besluit.add_run('Besluit:').bold = True

        doc.add_paragraph('Het raadsvoorstel als volgt te wijzigen:')

        doc.add_paragraph()

        # Wijzigingen
        for i, wijziging in enumerate(wijzigingen, 1):
            wijz_header = doc.add_paragraph()
            wijz_header.add_run(f'Wijziging {i}:').bold = True

            doc.add_paragraph()

            oor_label = doc.add_paragraph()
            oor_label.add_run('De tekst:').italic = True

            quote1 = doc.add_paragraph()
            quote1.paragraph_format.left_indent = Cm(1)
            quote1.add_run(f'"{wijziging.get("oorspronkelijk", "")}"')

            doc.add_paragraph()

            nieuw_label = doc.add_paragraph()
            nieuw_label.add_run('Te wijzigen in:').italic = True

            quote2 = doc.add_paragraph()
            quote2.paragraph_format.left_indent = Cm(1)
            quote2.add_run(f'"{wijziging.get("wordt", "")}"')

            doc.add_paragraph()

        # Toelichting
        if toelichting:
            doc.add_paragraph()
            toel = doc.add_paragraph()
            toel.add_run('Toelichting:').bold = True
            doc.add_paragraph(toelichting)

        doc.add_paragraph()
        doc.add_paragraph()

        # Ondertekening
        for indiener, partij in zip(indieners, partijen):
            doc.add_paragraph('_' * 40)
            doc.add_paragraph(f'{indiener} ({partij})')
            doc.add_paragraph()

        # Opslaan
        filename = self._generate_filename('amendement', titel)
        filepath = self.output_dir / filename
        doc.save(filepath)

        return filepath

    def get_generated_documents(self, doc_type: str = None) -> List[Dict]:
        """Lijst alle gegenereerde documenten."""
        documents = []
        patterns = ['*.docx', '*.md'] if doc_type is None else [f'{doc_type}_*.docx', f'{doc_type}_*.md']

        for pattern in patterns:
            for filepath in self.output_dir.glob(pattern):
                if doc_type and not filepath.stem.startswith(doc_type):
                    continue
                documents.append({
                    'filename': filepath.name,
                    'filepath': str(filepath),
                    'type': 'motie' if 'motie' in filepath.stem else 'amendement',
                    'created': filepath.stat().st_mtime
                })

        return sorted(documents, key=lambda x: x['created'], reverse=True)


# Singleton instance
_generator_instance = None


def get_document_generator() -> DocumentGenerator:
    """Get singleton DocumentGenerator instance."""
    global _generator_instance
    if _generator_instance is None:
        _generator_instance = DocumentGenerator()
    return _generator_instance


if __name__ == '__main__':
    # Test document generation
    generator = get_document_generator()

    # Test motie
    result = generator.generate_motie(
        titel="Meer groen in de binnenstad",
        indieners=["Jan Jansen", "Piet Pietersen"],
        partijen=["GroenLinks", "D66"],
        constateringen=[
            "dat de binnenstad van Baarn relatief weinig groen bevat",
            "dat meer groen bijdraagt aan een prettig leefklimaat"
        ],
        overwegingen=[
            "dat de gemeente streeft naar een duurzame leefomgeving",
            "dat burgers behoefte hebben aan meer groen in de openbare ruimte"
        ],
        verzoeken=[
            "om een plan op te stellen voor vergroening van de binnenstad",
            "om de raad hierover binnen 6 maanden te informeren"
        ],
        vergadering_datum="2024-03-15",
        agendapunt="7"
    )

    print(f"Generated: {result}")
